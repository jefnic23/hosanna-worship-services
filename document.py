import os
from itertools import chain

from docx import Document
from docx.shared import Inches, Pt

from utils import get_superscripts, pairwise


class WordDocument():
    def __init__(self, day):
        self._day = day
        self._document = Document()
        self._section = self._document.sections[0]

        self._section.left_margin = self._section.right_margin = self._section.top_margin = self._section.bottom_margin = Inches(0.5)
        
        if os.path.exists(f'services/{self._day}/first-reading.txt'):
            self._first_reading = open(f'services/{self._day}/first-reading.txt', 'r', encoding='utf-8').read()

        if os.path.exists(f'services/{self._day}/psalm.txt'):
            self._psalm = open(f'services/{self._day}/psalm.txt', 'r', encoding='utf-8').read()
        
        if os.path.exists(f'services/{self._day}/second-reading.txt'):
            self._second_reading = open(f'services/{self._day}/second-reading.txt', 'r', encoding='utf-8').read()

    def add_reading(self, reading):
        paragraph = self._document.add_paragraph()
        run = paragraph.add_run(reading.splitlines()[0] + "\n")
        run.font.bold = True
        run.font.size = Pt(14)

        for line in reading.splitlines()[1:-1]:
            superscripts = get_superscripts(line)
            if len(superscripts) > 0:
                index = pairwise(list(chain(*[[0], *[[s, e] for s, e in superscripts], [len(line)]])))
                for start, end in index:
                    run = paragraph.add_run(line[start:end])
                    run.font.superscript = True if (start, end) in superscripts else False
                    run.font.size = Pt(14)

    def save(self):
        if not os.path.exists(f'services/{self._day}'):
            os.makedirs(f'services/{self._day}')
        self._document.save(f'services/{self._day}/readings.docx')
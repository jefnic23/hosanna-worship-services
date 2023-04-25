import os
from datetime import date
from itertools import chain
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

from services.utils import clean_text, get_superscripts, lookahead, pairwise


class WordDocument():
    '''A class for creating a Word document.'''
    def __init__(
        self, 
        day: date,
        path: Path = Path('D:/Documents/Hosanna/services')
    ):
        self._day = day
        self._path = path
        self._document = Document()
        self._section = self._document.sections[0]
        self._section.left_margin =  \
        self._section.right_margin = \
        self._section.top_margin =   \
        self._section.bottom_margin = Inches(0.5)

    def add_reading(self, text: str) -> None:
        '''Add a reading to the document.'''
        paragraph = self._document.add_paragraph()
        run = paragraph.add_run(text.splitlines()[0] + "\n")
        run.font.bold = True
        run.font.size = Pt(14)
        
        reading = [clean_text(line) for line in text.splitlines()[1:]]
        for line, has_more in lookahead(reading):
            superscripts = get_superscripts(line)
            if len(superscripts) > 0:
                index = pairwise(
                    list(chain(*[[0], *[[s, e] for s, e in superscripts], [len(line)]]))
                )
                for start, end in index:
                    run = paragraph.add_run(line[start:end].strip())
                    run.font.superscript = True if (start, end) in superscripts else False
                    run.font.size = Pt(14)
            else:
                run = paragraph.add_run(line.strip())
                run.font.size = Pt(14)
                run.font.bold = False if has_more else True

            if has_more:
                run.add_break()


    def add_psalm(self, psalm: str) -> None:
        '''Add a psalm to the document.'''
        paragraph = self._document.add_paragraph()
        run = paragraph.add_run(psalm.splitlines()[0] + "\n")
        run.font.bold = True
        run.font.size = Pt(14)

        for i, line in enumerate(psalm.splitlines()[1:]):
            run = paragraph.add_run(line[0:2])
            run.font.superscript = True
            run.font.size = Pt(14)
            run.font.bold = False if i % 2 == 0 else True

            run = paragraph.add_run(line[2:].strip())
            run.font.superscript = False
            run.font.size = Pt(14)
            run.font.bold = False if i % 2 == 0 else True

            run.add_break()


    def save(self) -> None:
        '''Save the document to the services folder.'''        
        if not os.path.exists(f'{self._path}/{self._day}'):
            os.makedirs(f'{self._path}/{self._day}')
        self._document.save(f'{self._path}/{self._day}/{self._day}.docx')

        os.system(
            f'soffice --headless --invisible --convert-to pdf --outdir '
            f'{self._path}/{self._day} {self._path}/{self._day}/{self._day}.docx'
        )
        os.remove(f'{self._path}/{self._day}/{self._day}.docx')
        
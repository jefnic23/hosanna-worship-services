import os
import re
from datetime import date
from itertools import chain, pairwise
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph, Run

from services.settings import Settings
from services.utils import clean_text, get_superscripts, lookahead, split_formatted_text


class WordDocument:
    """A class for creating a Word document."""

    DEFAULT_FONTSIZE = 14

    def __init__(self, settings: Settings):
        self.day: date = date.today()
        self._path: Path = f"{settings.LOCAL_DIR}/services"
        self._document = Document()
        self._section = self._document.sections[0]
        self._section.left_margin = self._section.right_margin = (
            self._section.top_margin
        ) = self._section.bottom_margin = Inches(0.5)

        self.add_page_numbers()

    def add_rich_text(
        self,
        title: str | None,
        body: str | None,
        keep_together: bool = False,
        highlight_title: bool = False,
    ) -> None:
        """Add text to the document."""
        paragraph = self._document.add_paragraph()

        if title is not None:
            WordDocument._add_run(
                paragraph,
                f"{title}\n",
                bold=True,
                small_caps=True,
                highlight=highlight_title,
                underline=False if highlight_title else True
            )

        if body is None:
            return

        superscripts = re.findall(r"<sup>(.*?)</sup>", body)
        regular_text, bold_text, italic_text = split_formatted_text(
            body.replace("<sup>", "")
            .replace("</sup>", "")
            .replace("<div>", "")
            .replace("</div>", "")
            .replace("<br>", "")
        )
        bold_lines = [line[1] for line in bold_text]
        lines = [
            clean_text(line[1])
            for line in sorted(
                regular_text + bold_text + italic_text, key=lambda x: x[0]
            )
        ]
        for line, has_more in lookahead(lines):
            if not line:
                continue
            sups = get_superscripts(superscripts, line)
            if len(sups) > 0:
                index = pairwise(
                    list(chain(*[[0], *[[s, e] for s, e in sups], [len(line)]]))
                )
                for start, end in index:
                    run = WordDocument._add_run(
                        paragraph,
                        line[start:end],
                        superscript=True if (start, end) in sups else False,
                        bold=True if line in bold_lines else False,
                        color=(86, 86, 86) if (start, end) in sups else (0, 0, 0),
                    )
                for _ in range(len(sups)):
                    if superscripts:
                        superscripts.pop(0)
            else:
                run = WordDocument._add_run(
                    paragraph,
                    line,
                    bold=True if line in bold_lines else False,
                    color=(86, 86, 86) if line in sups else (0, 0, 0),
                )
            if has_more:
                run.add_break()

        paragraph.paragraph_format.keep_together = keep_together
        paragraph.paragraph_format.widow_control = True

    def add_page_numbers(self):
        """Add page numbers to the header of the document."""
        header = self._section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Create a new run and add the PAGE field to it
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")  # Create the fldChar element
        fldChar1.set(qn("w:fldCharType"), "begin")  # Set the fldCharType to 'begin'
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"  # Set the field code
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")  # Set the fldCharType to 'end'

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        # Add total number of pages field (optional)
        run = paragraph.add_run(" of ")
        run = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "NUMPAGES"  # Set the field code
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def save(self, filename: str, convert_to_pdf: bool = False) -> None:
        """Save the document to the services folder."""
        if not os.path.exists(f"{self._path}/{self.day}"):
            os.makedirs(f"{self._path}/{self.day}")
        self._document.save(f"{self._path}/{self.day}/{filename}.docx")

        if convert_to_pdf:
            self._convert_to_pdf(filename)
            os.remove(f"{self._path}/{self.day}/{filename}.docx")

    def _convert_to_pdf(self, filename: str) -> None:
        os.system(
            f"soffice --headless --invisible --convert-to pdf --outdir "
            f'{self._path}/{self.day} "{self._path}/{self.day}/{filename}.docx"'
        )

    def reset(self) -> None:
        self._document = Document()
        self._section = self._document.sections[0]
        self._section.left_margin = self._section.right_margin = (
            self._section.top_margin
        ) = self._section.bottom_margin = Inches(0.5)

        self.add_page_numbers()

    @staticmethod
    def _add_run(
        paragraph: Paragraph,
        text: str,
        size: int = DEFAULT_FONTSIZE,
        bold: bool = False,
        italic: bool = False,
        superscript: bool = False,
        color: tuple = (0, 0, 0),
        small_caps: bool = False,
        highlight: bool = False,
        underline: bool = False
    ) -> Run:
        """Add a run to a paragraph"""
        run = paragraph.add_run()
        run.text = text
        run.font.bold = bold
        run.font.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(*color)
        run.font.superscript = superscript
        run.font.small_caps = small_caps
        run.font.highlight_color = (
            WD_COLOR_INDEX.YELLOW if highlight else WD_COLOR_INDEX.AUTO
        )
        run.underline = underline
        return run
